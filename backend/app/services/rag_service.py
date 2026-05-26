import io
import logging

from langchain_core.documents import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_ollama import OllamaEmbeddings
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.config import settings as app_settings
from app.core.security import decrypt_api_key

logger = logging.getLogger(__name__)

# Embedding dimension per provider — must match model output
EMBEDDING_DIMENSIONS: dict[str, int] = {
    "openai": 1536,    # text-embedding-3-small
    "google": 768,     # models/text-embedding-004
    "ollama": 768,     # nomic-embed-text
    "anthropic": 768,  # falls back to nomic-embed-text
    "nvidia_nim": 768, # falls back to nomic-embed-text
}


def collection_name(user_id: str, doc_type: str) -> str:
    return f"{user_id}_{doc_type}"


def extract_text(content: bytes, filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        import fitz  # PyMuPDF
        doc = fitz.open(stream=content, filetype="pdf")
        return "\n".join(page.get_text() for page in doc)
    if lower.endswith(".docx"):
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)
    return content.decode("utf-8", errors="replace")


def chunk_text(text: str) -> list[str]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    return splitter.split_text(text)


def get_embedding_model(model_settings):
    provider = model_settings.provider
    if provider == "openai":
        api_key = decrypt_api_key(model_settings.api_key_enc, app_settings.APP_SECRET_KEY)
        return OpenAIEmbeddings(model="text-embedding-3-small", api_key=api_key)
    if provider == "google":
        api_key = decrypt_api_key(model_settings.api_key_enc, app_settings.APP_SECRET_KEY)
        return GoogleGenerativeAIEmbeddings(model="models/text-embedding-004", google_api_key=api_key)
    if provider == "ollama":
        return OllamaEmbeddings(model="nomic-embed-text", base_url=model_settings.ollama_url)
    # anthropic + nvidia_nim have no native embedding API — fall back to local nomic-embed-text
    return OllamaEmbeddings(model="nomic-embed-text")


def _psycopg_url() -> str:
    """Convert asyncpg URL to psycopg3 URL for langchain-postgres PGVector."""
    url = app_settings.DATABASE_URL
    if "+asyncpg" in url:
        return url.replace("+asyncpg", "+psycopg")
    if "postgresql://" in url and "+psycopg" not in url:
        logger.warning(
            "DATABASE_URL does not contain '+asyncpg' driver prefix. "
            "Assuming psycopg3 compatibility — verify your connection string."
        )
        return url.replace("postgresql://", "postgresql+psycopg://")
    return url


def get_vector_store(user_id: str, doc_type: str, embeddings, provider: str = "openai"):
    """Get or create PGVector store for a user+doc_type collection."""
    from langchain_postgres import PGVector

    table = collection_name(user_id, doc_type)
    return PGVector(
        connection_string=_psycopg_url(),
        collection_name=table,
        embedding=embeddings,
    )


def ingest_document(
    user_id: str,
    doc_type: str,
    text: str,
    metadata: dict,
    model_settings,
) -> int:
    """Chunk, embed, store. Returns chunk count."""
    chunks = chunk_text(text)
    embeddings = get_embedding_model(model_settings)
    docs = [
        Document(page_content=chunk, metadata={**metadata, "chunk_index": i})
        for i, chunk in enumerate(chunks)
    ]
    store = get_vector_store(user_id, doc_type, embeddings, provider=model_settings.provider)
    store.add_documents(docs)
    return len(docs)


def retrieve(
    user_id: str,
    doc_type: str,
    query: str,
    model_settings,
    k: int = 5,
) -> list[Document]:
    """Retrieve top-k relevant chunks."""
    embeddings = get_embedding_model(model_settings)
    store = get_vector_store(user_id, doc_type, embeddings, provider=model_settings.provider)
    return store.similarity_search(query, k=k)

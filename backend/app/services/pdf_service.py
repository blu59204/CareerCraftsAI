"""
ATS-safe PDF resume generation.

Three templates (classic, modern, technical), all single-column.
Compatible with Taleo, Workday, Greenhouse, iCIMS.
"""
import io
import logging
import re
from typing import Literal

from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

logger = logging.getLogger(__name__)

Template = Literal["classic", "modern", "technical"]

_EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
_PHONE_RE = re.compile(r"[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6}")
_LINKEDIN_RE = re.compile(r"linkedin\.com/in/", re.IGNORECASE)
_SECTION_NAMES_RE = re.compile(
    r"^(EXPERIENCE|EMPLOYMENT|WORK HISTORY|EDUCATION|SKILLS|TECHNOLOGIES|"
    r"CERTIFICATIONS|SUMMARY|OBJECTIVE|PROFILE|PROJECTS|PORTFOLIO|AWARDS|"
    r"PUBLICATIONS|LANGUAGES|INTERESTS|REFERENCES|ABOUT)",
    re.IGNORECASE,
)


def _escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _build_styles(template: Template) -> dict:
    base = getSampleStyleSheet()

    if template == "classic":
        name_style = ParagraphStyle(
            "Name", parent=base["Normal"], fontName="Times-Bold",
            fontSize=16, spaceAfter=4, alignment=TA_CENTER,
        )
        contact_style = ParagraphStyle(
            "Contact", parent=base["Normal"], fontName="Times-Roman",
            fontSize=10, spaceAfter=2, alignment=TA_CENTER,
        )
        section_style = ParagraphStyle(
            "Section", parent=base["Normal"], fontName="Times-Bold",
            fontSize=11, spaceBefore=10, spaceAfter=4, underline=True,
        )
        body_style = ParagraphStyle(
            "Body", parent=base["Normal"], fontName="Times-Roman",
            fontSize=10, leading=14, spaceAfter=2,
        )
        bullet_style = ParagraphStyle(
            "Bullet", parent=base["Normal"], fontName="Times-Roman",
            fontSize=10, leading=14, leftIndent=12, spaceAfter=2,
        )

    elif template == "technical":
        name_style = ParagraphStyle(
            "Name", parent=base["Normal"], fontName="Helvetica-Bold",
            fontSize=15, spaceAfter=4,
        )
        contact_style = ParagraphStyle(
            "Contact", parent=base["Normal"], fontName="Helvetica",
            fontSize=10, spaceAfter=2,
        )
        section_style = ParagraphStyle(
            "Section", parent=base["Normal"], fontName="Helvetica-Bold",
            fontSize=10, spaceBefore=8, spaceAfter=3, textTransform="uppercase",
        )
        body_style = ParagraphStyle(
            "Body", parent=base["Normal"], fontName="Helvetica",
            fontSize=10, leading=13, spaceAfter=1,
        )
        bullet_style = ParagraphStyle(
            "Bullet", parent=base["Normal"], fontName="Courier",
            fontSize=9, leading=13, leftIndent=12, spaceAfter=1,
        )

    else:  # modern (default)
        name_style = ParagraphStyle(
            "Name", parent=base["Normal"], fontName="Helvetica-Bold",
            fontSize=16, spaceAfter=4,
        )
        contact_style = ParagraphStyle(
            "Contact", parent=base["Normal"], fontName="Helvetica",
            fontSize=10, spaceAfter=2,
        )
        section_style = ParagraphStyle(
            "Section", parent=base["Normal"], fontName="Helvetica-Bold",
            fontSize=11, spaceBefore=10, spaceAfter=4,
        )
        body_style = ParagraphStyle(
            "Body", parent=base["Normal"], fontName="Helvetica",
            fontSize=10.5, leading=15, spaceAfter=2,
        )
        bullet_style = ParagraphStyle(
            "Bullet", parent=base["Normal"], fontName="Helvetica",
            fontSize=10.5, leading=15, leftIndent=12, spaceAfter=2,
        )

    return {
        "name": name_style,
        "contact": contact_style,
        "section": section_style,
        "body": body_style,
        "bullet": bullet_style,
    }


def _is_section_header(line: str) -> bool:
    stripped = line.strip()
    if not stripped or len(stripped) > 60:
        return False
    if stripped.isupper() and len(stripped) < 50:
        return True
    if _SECTION_NAMES_RE.match(stripped):
        return True
    if stripped.endswith(":") and len(stripped) < 40 and stripped[:-1].istitle():
        return True
    return False


def _is_contact_line(line: str) -> bool:
    return bool(
        _EMAIL_RE.search(line)
        or _PHONE_RE.search(line)
        or _LINKEDIN_RE.search(line)
    )


def _is_bullet_line(line: str) -> bool:
    stripped = line.strip()
    return bool(re.match(r"^[•\-\*]\s", stripped) or re.match(r"^\d+\.\s", stripped))


def generate_resume_pdf(
    resume_text: str,
    full_name: str = "",
    template: Template = "modern",
) -> bytes:
    if not resume_text or not resume_text.strip():
        raise ValueError("resume_text cannot be empty")

    styles = _build_styles(template)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    story = []
    lines = resume_text.strip().split("\n")
    name_written = False
    use_hr = template in ("modern", "technical")

    for line in lines:
        stripped = line.strip()

        if not stripped:
            story.append(Spacer(1, 4))
            continue

        if not name_written:
            display = _escape(full_name if full_name else stripped)
            story.append(Paragraph(display, styles["name"]))
            name_written = True
            if full_name and stripped != full_name:
                # The first line of the resume text is also contact/name info
                if _is_contact_line(stripped):
                    story.append(Paragraph(_escape(stripped), styles["contact"]))
            continue

        if _is_section_header(stripped):
            if use_hr:
                story.append(
                    HRFlowable(width="100%", thickness=0.5, color="black", spaceAfter=2)
                )
            story.append(Paragraph(_escape(stripped), styles["section"]))
            continue

        if _is_contact_line(stripped):
            story.append(Paragraph(_escape(stripped), styles["contact"]))
            continue

        if _is_bullet_line(stripped):
            story.append(Paragraph(_escape(stripped), styles["bullet"]))
            continue

        story.append(Paragraph(_escape(stripped), styles["body"]))

    doc.build(story)
    return buffer.getvalue()



def generate_resume_docx(resume_text: str, full_name: str = "") -> bytes:
    """Generate ATS-friendly DOCX from resume text.

    Uses Calibri 11pt, single column, standard margins.
    Compatible with Greenhouse, Workday, Taleo, iCIMS parsers.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    if not resume_text or not resume_text.strip():
        raise ValueError("resume_text cannot be empty")

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    lines = resume_text.strip().split("\n")
    name_written = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if not name_written:
            display = full_name if full_name else stripped
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(display)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Calibri"
            name_written = True
            if full_name and stripped != full_name and _is_contact_line(stripped):
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cp.add_run(stripped)
                cr.font.size = Pt(10)
                cr.font.name = "Calibri"
            continue

        if _is_section_header(stripped):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            # Add bottom border via paragraph formatting
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            continue

        if _is_contact_line(stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            continue

        if _is_bullet_line(stripped):
            p = doc.add_paragraph(style="List Bullet")
            # Remove bullet prefix for clean rendering
            clean = re.sub(r"^[•\-\*]\s*", "", stripped)
            clean = re.sub(r"^\d+\.\s*", "", clean)
            run = p.add_run(clean)
            run.font.size = Pt(10.5)
            run.font.name = "Calibri"
            continue

        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
